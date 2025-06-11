import os
import re
from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.responses import JSONResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
from docx import Document
import pytesseract
from pdf2image import convert_from_bytes
from io import BytesIO
from docxtpl import DocxTemplate

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

def clean_text_single_line(text: str) -> str:
    text = text.replace('\n', ' ').replace('\r', ' ')
    text = re.sub(r'\s+', ' ', text)
    symbols = ['/', '|', '-', '•']
    for sym in symbols:
        text = re.sub(rf'\s*{re.escape(sym)}\s*', f' {sym} ', text)
    return text.strip()

async def extract_text_from_docx(file: UploadFile) -> str:
    content = await file.read()
    file_stream = BytesIO(content)
    doc = Document(file_stream)
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            row_text = ' '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                full_text.append(row_text)
    return clean_text_single_line(' '.join(full_text))

async def extract_text_from_pdf(file: UploadFile) -> str:
    content = await file.read()
    images = convert_from_bytes(content)
    extracted_text = []
    for image in images:
        custom_config = r'--psm 4'
        page_text = pytesseract.image_to_string(image, config=custom_config)
        if page_text.strip():
            extracted_text.append(page_text.strip())
    return clean_text_single_line(' '.join(extracted_text))

def parse_extracted_text(text: str) -> dict:
    result = {
        "FULL_NAME": "",
        "NATIONALITY": "",
        "LOCATION": "",
        "LANGUAGES": "",
        "EDUCATION": "",
        "EMPLOYMENT_HISTORY": ""
    }

    # FULL_NAME
    words = text.strip().split()
    if len(words) >= 2:
        result["FULL_NAME"] = f"{words[0].capitalize()} {words[1].capitalize()}"

    # NATIONALITY
    match_nationality = re.search(r'Nationality[:\-]?\s*([A-Za-z]+)', text, re.IGNORECASE)
    if match_nationality:
        result["NATIONALITY"] = match_nationality.group(1).capitalize()

    # LOCATION
    match_location = re.search(r'Location[:\-]?\s*([A-Za-z\s]+)', text, re.IGNORECASE)
    if match_location:
        city = match_location.group(1).split(',')[0].strip()
        result["LOCATION"] = city.capitalize()

    # LANGUAGES
    match_languages = re.search(r'Languages[:\-]?\s*([^\n]+)', text, re.IGNORECASE)
    if match_languages:
        raw_langs = match_languages.group(1)
        langs = re.findall(r'\b[A-Za-z]+\b', raw_langs)
        if langs:
            result["LANGUAGES"] = ", ".join([lang.capitalize() for lang in langs])

    # EDUCATION
    education_entries = []
    matches = re.findall(r'(\d{2,4})[^\d]+([^\n:]+)[:,\-]?\s*([^\n]+)', text, re.IGNORECASE)
    for match in matches:
        year = match[0]
        university = match[1].strip().title() + ":"
        degree = match[2].strip()
        education_entries.append(f"{year}    {university}\n    {degree}")
    result["EDUCATION"] = "\n\n".join(education_entries) if education_entries else ""

    # EMPLOYMENT_HISTORY
    employment_entries = []
    emp_matches = re.findall(r'(\d{1,2}[\/\-]\d{2,4}|\d{4})[^\d]+([^\n]+)', text)
    for match in emp_matches:
        date_raw = match[0]
        company_position = match[1].strip()

        # Date formatting
        try:
            if '/' in date_raw or '-' in date_raw:
                parts = re.split(r'[-/]', date_raw)
                if len(parts) >= 2:
                    start = format_date(parts[0])
                    end = format_date(parts[1])
                    date = f"{start} – {end}"
                else:
                    date = format_date(parts[0])
            else:
                date = date_raw.strip()
        except:
            date = date_raw.strip()

        # Company & Position formatting
        parts = company_position.split(',')
        company = parts[0].strip().title()
        position = ', '.join(parts[1:]).strip().title() if len(parts) > 1 else ''

        # Append
        employment_entries.append(f"{date}    {company}\n    {position}")

    result["EMPLOYMENT_HISTORY"] = "\n\n".join(employment_entries) if employment_entries else ""

    return result

def format_date(raw_date):
    """
    Format date to Month Year or Year only based on the input.
    """
    raw_date = raw_date.strip().replace("‘", "").replace("’", "")
    if re.match(r'\d{4}$', raw_date):
        return raw_date
    elif re.match(r'\d{1,2}/\d{2,4}$', raw_date):
        month_num, year = raw_date.split('/')
        month = month_name(int(month_num))
        year = "20" + year if len(year) == 2 else year
        return f"{month} {year}"
    else:
        return raw_date

def month_name(month_num):
    months = ["Jan", "Feb", "Mar", "Apr", "May", "Jun",
              "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"]
    try:
        return months[int(month_num) - 1]
    except:
        return ""

def generate_cv(context: dict) -> str:
    try:
        template_path = os.path.join("templates", "uk_danos_compliance.docx")
        doc = DocxTemplate(template_path)
        doc.render(context)
        full_name = context.get("FULL_NAME", "Candidate").strip()
        safe_name = re.sub(r'[^\w\s-]', '', full_name).replace(' ', '_')
        output_filename = f"{safe_name}.docx"
        output_path = os.path.join("generated_cvs", output_filename)
        os.makedirs("generated_cvs", exist_ok=True)
        doc.save(output_path)
        return output_path
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"CV generation failed: {str(e)}")

@app.post("/upload/")
async def upload_cv(file: UploadFile = File(...)):
    try:
        filename = file.filename.lower()
        if filename.endswith(".docx"):
            text = await extract_text_from_docx(file)
        elif filename.endswith(".pdf"):
            text = await extract_text_from_pdf(file)
        else:
            raise HTTPException(status_code=400, detail="Unsupported file format. Please upload .docx or .pdf")
        return JSONResponse({"extracted_text": text})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Internal server error: {str(e)}")

@app.post("/generate-cv/")
async def generate_cv_endpoint(context: dict):
    extracted_text = context.get("extracted_text", "")
    parsed_fields = parse_extracted_text(extracted_text)
    context.update(parsed_fields)
    output_path = generate_cv(context)
    return JSONResponse({"message": "CV generated successfully!", "output_file": output_path})

@app.get("/download-cv/")
async def download_cv(filename: str):
    file_path = os.path.join("generated_cvs", filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")
    return FileResponse(
        path=file_path,
        filename=filename,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@app.get("/")
async def root():
    return {"message": "Welcome to CVFormatic AI Backend"}
